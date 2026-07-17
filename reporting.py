from __future__ import annotations

import io
import math
from datetime import datetime
from typing import Any, Iterable

import pandas as pd
import xlsxwriter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageOps


COR_PRIMARIA = "2471A3"
COR_SECUNDARIA = "D6EAF8"
COR_CABECALHO = "1F4E78"
COR_CLARA = "F4F6F7"
COR_BORDA = "B7C9D6"


def _texto(valor: Any, padrao: str = "-") -> str:
    """Converte valores nulos em texto legível para relatórios."""
    if valor is None:
        return padrao
    try:
        if pd.isna(valor):
            return padrao
    except (TypeError, ValueError):
        pass
    conteudo = str(valor).strip()
    return conteudo or padrao


def _numero(valor: Any, casas: int = 2, sufixo: str = "") -> str:
    """Formata valores numéricos com segurança."""
    if valor is None:
        return "-"
    try:
        numero = float(valor)
    except (TypeError, ValueError):
        return _texto(valor)
    if not math.isfinite(numero):
        return "-"
    return f"{numero:.{casas}f}{sufixo}"


def _normalizar_valor_excel(valor: Any) -> Any:
    """Converte tipos do Pandas em valores aceitos pelo XlsxWriter."""
    if valor is None:
        return None
    try:
        if pd.isna(valor):
            return None
    except (TypeError, ValueError):
        pass
    if isinstance(valor, (pd.Timestamp, datetime)):
        return valor.to_pydatetime() if isinstance(valor, pd.Timestamp) else valor
    if hasattr(valor, "item"):
        try:
            return valor.item()
        except (AttributeError, ValueError):
            pass
    return valor


def _aplicar_sombreamento(celula: Any, cor: str) -> None:
    """Aplica preenchimento de fundo a uma célula do Word."""
    propriedades = celula._tc.get_or_add_tcPr()
    elemento = propriedades.find(qn("w:shd"))
    if elemento is None:
        elemento = OxmlElement("w:shd")
        propriedades.append(elemento)
    elemento.set(qn("w:fill"), cor)


def _aplicar_bordas_tabela(tabela: Any, cor: str = COR_BORDA) -> None:
    """Aplica bordas discretas à tabela do Word."""
    propriedades = tabela._tbl.tblPr
    bordas = propriedades.first_child_found_in("w:tblBorders")
    if bordas is None:
        bordas = OxmlElement("w:tblBorders")
        propriedades.append(bordas)
    for nome in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elemento = bordas.find(qn(f"w:{nome}"))
        if elemento is None:
            elemento = OxmlElement(f"w:{nome}")
            bordas.append(elemento)
        elemento.set(qn("w:val"), "single")
        elemento.set(qn("w:sz"), "4")
        elemento.set(qn("w:color"), cor)


def _definir_repeticao_cabecalho(linha: Any) -> None:
    """Marca a primeira linha para repetição em tabelas longas no Word."""
    propriedades = linha._tr.get_or_add_trPr()
    repetir = OxmlElement("w:tblHeader")
    repetir.set(qn("w:val"), "true")
    propriedades.append(repetir)


def _adicionar_tabela_chave_valor(
    documento: Document,
    pares: Iterable[tuple[str, Any]],
    colunas: int = 2,
) -> Any:
    """Cria uma tabela compacta de metadados em pares rótulo/valor."""
    itens = list(pares)
    linhas = math.ceil(len(itens) / colunas)
    tabela = documento.add_table(rows=linhas, cols=colunas * 2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = True
    _aplicar_bordas_tabela(tabela)

    indice = 0
    for linha in tabela.rows:
        for grupo in range(colunas):
            celula_rotulo = linha.cells[grupo * 2]
            celula_valor = linha.cells[grupo * 2 + 1]
            celula_rotulo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            celula_valor.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if indice >= len(itens):
                celula_rotulo.text = ""
                celula_valor.text = ""
                continue
            rotulo, valor = itens[indice]
            celula_rotulo.text = str(rotulo)
            celula_valor.text = _texto(valor)
            _aplicar_sombreamento(celula_rotulo, COR_SECUNDARIA)
            for paragrafo in celula_rotulo.paragraphs:
                for run in paragrafo.runs:
                    run.bold = True
                    run.font.size = Pt(8.5)
            for paragrafo in celula_valor.paragraphs:
                for run in paragrafo.runs:
                    run.font.size = Pt(8.5)
            indice += 1
    return tabela


def _adicionar_dataframe_word(
    documento: Document,
    dataframe: pd.DataFrame,
    colunas: list[tuple[str, str]],
    formatos: dict[str, str] | None = None,
    limite_linhas: int | None = None,
) -> Any:
    """Inclui um DataFrame como tabela paginável no Word."""
    formatos = formatos or {}
    dados = dataframe.copy()
    if limite_linhas is not None:
        dados = dados.head(limite_linhas)

    tabela = documento.add_table(rows=1, cols=len(colunas))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = True
    _aplicar_bordas_tabela(tabela)
    _definir_repeticao_cabecalho(tabela.rows[0])

    for indice, (_, titulo) in enumerate(colunas):
        celula = tabela.rows[0].cells[indice]
        celula.text = titulo
        _aplicar_sombreamento(celula, COR_CABECALHO)
        for paragrafo in celula.paragraphs:
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragrafo.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)

    if dados.empty:
        linha = tabela.add_row()
        linha.cells[0].text = "Sem registros."
        if len(colunas) > 1:
            linha.cells[0].merge(linha.cells[-1])
        return tabela

    for _, registro in dados.iterrows():
        linha = tabela.add_row()
        for indice, (campo, _) in enumerate(colunas):
            valor = registro.get(campo)
            formato = formatos.get(campo)
            if formato and valor is not None and not pd.isna(valor):
                try:
                    valor_texto = formato.format(float(valor))
                except (TypeError, ValueError):
                    valor_texto = _texto(valor)
            else:
                valor_texto = _texto(valor)
            celula = linha.cells[indice]
            celula.text = valor_texto
            celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragrafo in celula.paragraphs:
                for run in paragrafo.runs:
                    run.font.size = Pt(7.8)
    return tabela


def _preparar_foto_para_relatorio(conteudo: bytes, largura_maxima_px: int = 1600) -> io.BytesIO:
    """Corrige orientação e reduz fotografias antes de inseri-las no relatório."""
    origem = io.BytesIO(bytes(conteudo))
    with Image.open(origem) as imagem_aberta:
        imagem = ImageOps.exif_transpose(imagem_aberta).convert("RGB")
        if imagem.width > largura_maxima_px:
            fator = largura_maxima_px / imagem.width
            imagem = imagem.resize(
                (largura_maxima_px, max(1, int(imagem.height * fator))),
                Image.Resampling.LANCZOS,
            )
        destino = io.BytesIO()
        imagem.save(destino, format="JPEG", quality=86, optimize=True)
        destino.seek(0)
        return destino


def _adicionar_imagem_word(
    documento: Document,
    conteudo: bytes | None,
    legenda: str,
    largura_cm: float = 16.5,
) -> None:
    """Inclui uma imagem centralizada com legenda no Word."""
    if not conteudo:
        return
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run()
    run.add_picture(io.BytesIO(conteudo), width=Cm(largura_cm))
    legenda_paragrafo = documento.add_paragraph(legenda)
    legenda_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legenda_paragrafo.style = documento.styles["Caption"]


def gerar_relatorio_word(
    dados: dict[str, Any],
    imagem_perfil: bytes | None = None,
    imagem_construcao: bytes | None = None,
    imagem_desenvolvimento: bytes | None = None,
    incluir_fotos: bool = True,
    limite_fotos: int = 30,
) -> bytes:
    """Gera relatório técnico editável em formato DOCX."""
    projeto = dados.get("projeto") or {}
    sondagem = dados.get("sondagem") or {}
    camadas: pd.DataFrame = dados.get("camadas", pd.DataFrame())
    coletas: pd.DataFrame = dados.get("coletas", pd.DataFrame())
    voc: pd.DataFrame = dados.get("voc", pd.DataFrame())
    leituras_na: pd.DataFrame = dados.get("leituras_na", pd.DataFrame())
    poco = dados.get("poco") or {}
    intervalos: pd.DataFrame = dados.get("intervalos_construtivos", pd.DataFrame())
    desenvolvimento = dados.get("desenvolvimento") or {}
    leituras_desenvolvimento: pd.DataFrame = dados.get(
        "leituras_desenvolvimento", pd.DataFrame()
    )
    fotos: list[dict[str, Any]] = list(dados.get("fotos") or [])

    documento = Document()
    secao = documento.sections[0]
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.5)
    secao.left_margin = Cm(1.55)
    secao.right_margin = Cm(1.55)

    estilos = documento.styles
    estilos["Normal"].font.name = "Aptos"
    estilos["Normal"].font.size = Pt(9)
    for nome, tamanho, cor in (
        ("Title", 20, COR_CABECALHO),
        ("Heading 1", 15, COR_CABECALHO),
        ("Heading 2", 12, COR_PRIMARIA),
    ):
        estilo = estilos[nome]
        estilo.font.name = "Aptos Display" if nome == "Title" else "Aptos"
        estilo.font.size = Pt(tamanho)
        estilo.font.color.rgb = RGBColor.from_string(cor)

    titulo = documento.add_paragraph(style="Title")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.add_run("RELATÓRIO DE SONDAGEM E INSTALAÇÃO DE POÇO DE MONITORAMENTO")
    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(
        f"{_texto(projeto.get('nome'))} — {_texto(sondagem.get('nome_furo'))}"
    )
    run.bold = True
    run.font.size = Pt(12)
    documento.add_paragraph()

    documento.add_heading("1. Identificação", level=1)
    _adicionar_tabela_chave_valor(
        documento,
        [
            ("Projeto", projeto.get("nome")),
            ("Cliente", projeto.get("cliente")),
            ("Local", projeto.get("localizacao")),
            ("Sondagem", sondagem.get("nome_furo")),
            ("Status", sondagem.get("status")),
            ("Data de planejamento", sondagem.get("data")),
            ("Data de início", sondagem.get("data_inicio")),
            ("Data de conclusão", sondagem.get("data_conclusao")),
            ("Responsável técnico", projeto.get("responsavel_tecnico")),
            ("Registro profissional", projeto.get("registro_profissional")),
        ],
    )

    documento.add_heading("2. Localização e referência espacial", level=1)
    _adicionar_tabela_chave_valor(
        documento,
        [
            ("Sistema de entrada", f"EPSG:{_texto(sondagem.get('crs_entrada'))}"),
            ("Coordenada X", _numero(sondagem.get("coordenada_x"), 3)),
            ("Coordenada Y", _numero(sondagem.get("coordenada_y"), 3)),
            ("Latitude SIRGAS 2000", _numero(sondagem.get("latitude"), 8)),
            ("Longitude SIRGAS 2000", _numero(sondagem.get("longitude"), 8)),
            ("Origem da coordenada", sondagem.get("origem_coordenada")),
            ("Precisão GPS estimada", _numero(sondagem.get("precisao_gps_m"), 1, " m")),
            ("Data da captura GPS", sondagem.get("data_captura_gps")),
            ("Cota do terreno", _numero(sondagem.get("altitude"), 3, " m")),
        ],
    )

    documento.add_heading("3. Execução da sondagem", level=1)
    _adicionar_tabela_chave_valor(
        documento,
        [
            ("Método de perfuração", sondagem.get("metodo_perfuracao")),
            ("Equipamento", sondagem.get("equipamento")),
            ("Empresa executora", sondagem.get("empresa_executora")),
            ("Responsável de campo", sondagem.get("responsavel_campo")),
            ("Profundidade planejada", _numero(sondagem.get("profundidade_planejada"), 3, " m")),
            ("Profundidade executada", _numero(sondagem.get("profundidade_total"), 3, " m")),
            ("NA estático adotado", _numero(sondagem.get("nivel_agua_estatico"), 3, " m")),
            ("Observações", sondagem.get("observacoes_gerais")),
        ],
    )

    documento.add_heading("4. Perfil litológico", level=1)
    _adicionar_imagem_word(
        documento,
        imagem_perfil,
        "Figura 1 — Perfil litológico, condição hídrica, amostras e VOC.",
    )
    _adicionar_dataframe_word(
        documento,
        camadas,
        [
            ("profundidade_inicial", "De (m)"),
            ("profundidade_final", "Até (m)"),
            ("classificacao", "Classificação"),
            ("descricao_tatil_visual", "Descrição tátil-visual"),
            ("tipo_aquifero", "Unidade hidroestratigráfica"),
            ("zona_hidrica", "Condição hídrica"),
        ],
        formatos={
            "profundidade_inicial": "{:.3f}",
            "profundidade_final": "{:.3f}",
        },
    )

    documento.add_heading("5. Amostras, VOC e níveis d'água", level=1)
    documento.add_heading("5.1 Pontos de coleta", level=2)
    _adicionar_dataframe_word(
        documento,
        coletas,
        [("profundidade_coleta", "Profundidade da coleta (m)")],
        formatos={"profundidade_coleta": "{:.3f}"},
    )
    documento.add_heading("5.2 Medições de VOC", level=2)
    _adicionar_dataframe_word(
        documento,
        voc,
        [
            ("profundidade", "Profundidade (m)"),
            ("concentracao", "Concentração (mg/L ou ppm)"),
        ],
        formatos={"profundidade": "{:.3f}", "concentracao": "{:.6g}"},
    )
    documento.add_heading("5.3 Histórico de nível d'água", level=2)
    _adicionar_dataframe_word(
        documento,
        leituras_na,
        [
            ("data_hora", "Data e hora"),
            ("tipo", "Tipo"),
            ("profundidade_m", "Profundidade (m)"),
            ("usar_como_estatico", "NA adotado"),
            ("observacoes", "Observações"),
        ],
        formatos={"profundidade_m": "{:.3f}"},
    )

    documento.add_heading("6. Perfil construtivo do poço de monitoramento", level=1)
    if poco:
        _adicionar_tabela_chave_valor(
            documento,
            [
                ("Data de instalação", poco.get("data_instalacao")),
                ("Profundidade do poço", _numero(poco.get("profundidade_poco"), 3, " m")),
                ("Diâmetro da perfuração", _numero(poco.get("diametro_perfuracao_mm"), 1, " mm")),
                ("Diâmetro do revestimento", _numero(poco.get("diametro_revestimento_mm"), 1, " mm")),
                ("Material do revestimento", poco.get("material_revestimento")),
                ("Fabricante / modelo", poco.get("fabricante_modelo")),
                ("Cota da boca do tubo", _numero(poco.get("cota_boca_tubo"), 3, " m")),
                ("Altura da boca do tubo", _numero(poco.get("altura_boca_tubo_m"), 3, " m")),
                ("Proteção superficial", poco.get("tipo_protecao_superficial")),
                ("Câmara de calçada", "Sim" if poco.get("camara_calcada") else "Não"),
                ("Tampa", poco.get("tampa")),
                ("Responsável pela instalação", poco.get("responsavel_instalacao")),
                ("Observações", poco.get("observacoes")),
            ],
        )
        _adicionar_imagem_word(
            documento,
            imagem_construcao,
            "Figura 2 — Perfil construtivo do poço de monitoramento.",
        )
        _adicionar_dataframe_word(
            documento,
            intervalos,
            [
                ("componente", "Componente"),
                ("profundidade_inicial", "De (m)"),
                ("profundidade_final", "Até (m)"),
                ("material", "Material"),
                ("especificacao", "Especificação"),
                ("diametro_mm", "Diâmetro (mm)"),
                ("abertura_ranhura_mm", "Ranhura (mm)"),
                ("granulometria", "Granulometria"),
            ],
            formatos={
                "profundidade_inicial": "{:.3f}",
                "profundidade_final": "{:.3f}",
                "diametro_mm": "{:.2f}",
                "abertura_ranhura_mm": "{:.3f}",
            },
        )
    else:
        documento.add_paragraph("O perfil construtivo não foi cadastrado.")

    documento.add_heading("7. Desenvolvimento do poço", level=1)
    if desenvolvimento:
        if desenvolvimento.get("realizado"):
            _adicionar_tabela_chave_valor(
                documento,
                [
                    ("Realizado", "Sim"),
                    ("Data", desenvolvimento.get("data")),
                    ("Método", desenvolvimento.get("metodo")),
                    ("Duração", _numero(desenvolvimento.get("duracao_min"), 1, " min")),
                    ("Profundidade do equipamento", _numero(desenvolvimento.get("profundidade_equipamento_m"), 2, " m")),
                    ("NA antes", _numero(desenvolvimento.get("na_antes_m"), 3, " m")),
                    ("NA depois", _numero(desenvolvimento.get("na_depois_m"), 3, " m")),
                    ("Vazão", _numero(desenvolvimento.get("vazao_l_min"), 2, " L/min")),
                    ("Volume retirado", _numero(desenvolvimento.get("volume_retirado_l"), 1, " L")),
                    ("Turbidez inicial", _numero(desenvolvimento.get("turbidez_inicial_ntu"), 1, " NTU")),
                    ("Turbidez final", _numero(desenvolvimento.get("turbidez_final_ntu"), 1, " NTU")),
                    ("pH final", _numero(desenvolvimento.get("ph_final"), 2)),
                    ("Condutividade final", _numero(desenvolvimento.get("condutividade_final_us_cm"), 1, " µS/cm")),
                    ("Temperatura final", _numero(desenvolvimento.get("temperatura_final_c"), 1, " °C")),
                    ("Responsável", desenvolvimento.get("responsavel")),
                    ("Observações", desenvolvimento.get("observacoes")),
                ],
            )
            _adicionar_dataframe_word(
                documento,
                leituras_desenvolvimento,
                [
                    ("tempo_min", "Tempo (min)"),
                    ("nivel_agua_m", "NA (m)"),
                    ("vazao_l_min", "Vazão (L/min)"),
                    ("turbidez_ntu", "Turbidez (NTU)"),
                    ("ph", "pH"),
                    ("condutividade_us_cm", "Condutividade (µS/cm)"),
                    ("temperatura_c", "Temperatura (°C)"),
                    ("observacoes", "Observações"),
                ],
                formatos={
                    "tempo_min": "{:.1f}",
                    "nivel_agua_m": "{:.3f}",
                    "vazao_l_min": "{:.2f}",
                    "turbidez_ntu": "{:.1f}",
                    "ph": "{:.2f}",
                    "condutividade_us_cm": "{:.1f}",
                    "temperatura_c": "{:.1f}",
                },
            )
            _adicionar_imagem_word(
                documento,
                imagem_desenvolvimento,
                "Figura 3 — Evolução das leituras durante o desenvolvimento do poço.",
                largura_cm=16.0,
            )
        else:
            documento.add_paragraph(
                "Desenvolvimento não realizado. Motivo: "
                + _texto(desenvolvimento.get("motivo_nao_realizado"))
            )
    else:
        documento.add_paragraph("O desenvolvimento do poço não foi informado.")

    documento.add_heading("8. Registro fotográfico", level=1)
    if incluir_fotos and fotos:
        for indice, foto in enumerate(fotos[: max(0, int(limite_fotos))], start=1):
            try:
                imagem = _preparar_foto_para_relatorio(foto["conteudo"])
                paragrafo = documento.add_paragraph()
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragrafo.add_run()
                run.add_picture(imagem, width=Cm(14.5))
                legenda = (
                    f"Foto {indice} — {_texto(foto.get('categoria'))}. "
                    f"{_texto(foto.get('legenda'), '')}"
                ).strip()
                if foto.get("profundidade_m") is not None:
                    legenda += f" Profundidade associada: {_numero(foto.get('profundidade_m'), 2, ' m')}."
                legenda_paragrafo = documento.add_paragraph(legenda)
                legenda_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                legenda_paragrafo.style = documento.styles["Caption"]
            except Exception as erro:
                documento.add_paragraph(
                    f"Não foi possível inserir a fotografia {_texto(foto.get('nome_arquivo'))}: {erro}"
                )
        if len(fotos) > limite_fotos:
            documento.add_paragraph(
                f"O banco possui {len(fotos)} fotografias; as primeiras {limite_fotos} foram incluídas neste relatório."
            )
    else:
        documento.add_paragraph("Não há fotografias selecionadas para este relatório.")

    documento.add_heading("9. Observações e validação", level=1)
    documento.add_paragraph(
        "As informações deste documento foram geradas a partir dos dados cadastrados no diário digital. "
        "A interpretação hidrogeológica, a simbologia e o atendimento às normas aplicáveis devem ser "
        "conferidos e aprovados pelo responsável técnico antes da emissão oficial."
    )

    rodape = secao.footer.paragraphs[0]
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape_run = rodape.add_run(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} — Diário de Sondagem Hidrogeológica"
    )
    rodape_run.font.size = Pt(8)
    rodape_run.font.color.rgb = RGBColor(90, 90, 90)

    memoria = io.BytesIO()
    documento.save(memoria)
    return memoria.getvalue()


def _escrever_tabela_excel(
    planilha: Any,
    workbook: Any,
    dataframe: pd.DataFrame,
    colunas: list[tuple[str, str]],
    linha_inicial: int = 0,
) -> int:
    """Escreve uma tabela formatada e devolve a próxima linha livre."""
    formato_cabecalho = workbook.add_format(
        {
            "bold": True,
            "font_color": "#FFFFFF",
            "bg_color": "#1F4E78",
            "border": 1,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
        }
    )
    formato_texto = workbook.add_format(
        {"border": 1, "valign": "top", "text_wrap": True}
    )
    formato_numero = workbook.add_format(
        {"border": 1, "valign": "top", "num_format": "0.000"}
    )

    for coluna, (_, titulo) in enumerate(colunas):
        planilha.write(linha_inicial, coluna, titulo, formato_cabecalho)
    linha = linha_inicial + 1
    if dataframe.empty:
        planilha.write(linha, 0, "Sem registros", formato_texto)
        return linha + 2

    for _, registro in dataframe.iterrows():
        for coluna, (campo, _) in enumerate(colunas):
            valor = _normalizar_valor_excel(registro.get(campo))
            if isinstance(valor, (int, float)) and not isinstance(valor, bool):
                planilha.write_number(linha, coluna, float(valor), formato_numero)
            else:
                planilha.write(linha, coluna, valor, formato_texto)
        linha += 1
    return linha + 1


def gerar_relatorio_excel(
    dados: dict[str, Any],
    imagem_perfil: bytes | None = None,
    imagem_construcao: bytes | None = None,
    imagem_desenvolvimento: bytes | None = None,
    incluir_fotos: bool = True,
    limite_fotos: int = 30,
) -> bytes:
    """Gera anexo técnico estruturado em formato XLSX."""
    projeto = dados.get("projeto") or {}
    sondagem = dados.get("sondagem") or {}
    camadas: pd.DataFrame = dados.get("camadas", pd.DataFrame())
    coletas: pd.DataFrame = dados.get("coletas", pd.DataFrame())
    voc: pd.DataFrame = dados.get("voc", pd.DataFrame())
    leituras_na: pd.DataFrame = dados.get("leituras_na", pd.DataFrame())
    poco = dados.get("poco") or {}
    intervalos: pd.DataFrame = dados.get("intervalos_construtivos", pd.DataFrame())
    desenvolvimento = dados.get("desenvolvimento") or {}
    leituras_desenvolvimento: pd.DataFrame = dados.get(
        "leituras_desenvolvimento", pd.DataFrame()
    )
    fotos: list[dict[str, Any]] = list(dados.get("fotos") or [])

    memoria = io.BytesIO()
    workbook = xlsxwriter.Workbook(memoria, {"in_memory": True})
    workbook.set_properties(
        {
            "title": f"Relatório da sondagem {_texto(sondagem.get('nome_furo'))}",
            "subject": "Sondagem hidrogeológica e poço de monitoramento",
            "author": "Diário de Sondagem Hidrogeológica",
            "comments": "Arquivo gerado automaticamente a partir do banco SQLite.",
        }
    )

    formato_titulo = workbook.add_format(
        {
            "bold": True,
            "font_size": 18,
            "font_color": "#1F4E78",
            "align": "center",
            "valign": "vcenter",
        }
    )
    formato_secao = workbook.add_format(
        {
            "bold": True,
            "font_size": 12,
            "font_color": "#FFFFFF",
            "bg_color": "#2471A3",
            "border": 1,
        }
    )
    formato_rotulo = workbook.add_format(
        {"bold": True, "bg_color": "#D6EAF8", "border": 1, "text_wrap": True}
    )
    formato_valor = workbook.add_format({"border": 1, "text_wrap": True})

    resumo = workbook.add_worksheet("Resumo")
    resumo.hide_gridlines(2)
    resumo.set_column("A:A", 25)
    resumo.set_column("B:B", 30)
    resumo.set_column("C:C", 25)
    resumo.set_column("D:D", 30)
    resumo.merge_range("A1:D2", "RELATÓRIO DE SONDAGEM E POÇO DE MONITORAMENTO", formato_titulo)
    resumo.write("A4", "Identificação", formato_secao)
    resumo.merge_range("A4:D4", "Identificação", formato_secao)

    pares = [
        ("Projeto", projeto.get("nome")),
        ("Cliente", projeto.get("cliente")),
        ("Local", projeto.get("localizacao")),
        ("Sondagem", sondagem.get("nome_furo")),
        ("Status", sondagem.get("status")),
        ("EPSG", sondagem.get("crs_entrada")),
        ("Coordenada X", sondagem.get("coordenada_x")),
        ("Coordenada Y", sondagem.get("coordenada_y")),
        ("Latitude", sondagem.get("latitude")),
        ("Longitude", sondagem.get("longitude")),
        ("Origem da coordenada", sondagem.get("origem_coordenada")),
        ("Precisão GPS (m)", sondagem.get("precisao_gps_m")),
        ("Cota do terreno (m)", sondagem.get("altitude")),
        ("Profundidade planejada (m)", sondagem.get("profundidade_planejada")),
        ("Profundidade executada (m)", sondagem.get("profundidade_total")),
        ("NA estático (m)", sondagem.get("nivel_agua_estatico")),
        ("Método de perfuração", sondagem.get("metodo_perfuracao")),
        ("Equipamento", sondagem.get("equipamento")),
        ("Empresa executora", sondagem.get("empresa_executora")),
        ("Responsável de campo", sondagem.get("responsavel_campo")),
    ]
    linha = 5
    for indice in range(0, len(pares), 2):
        for grupo, item in enumerate(pares[indice : indice + 2]):
            coluna = grupo * 2
            rotulo, valor = item
            resumo.write(linha, coluna, rotulo, formato_rotulo)
            resumo.write(linha, coluna + 1, _normalizar_valor_excel(valor), formato_valor)
        linha += 1

    if poco:
        linha += 1
        resumo.merge_range(linha, 0, linha, 3, "Dados gerais do poço", formato_secao)
        linha += 1
        pares_poco = [
            ("Data de instalação", poco.get("data_instalacao")),
            ("Profundidade do poço (m)", poco.get("profundidade_poco")),
            ("Diâmetro da perfuração (mm)", poco.get("diametro_perfuracao_mm")),
            ("Diâmetro do revestimento (mm)", poco.get("diametro_revestimento_mm")),
            ("Material do revestimento", poco.get("material_revestimento")),
            ("Proteção superficial", poco.get("tipo_protecao_superficial")),
            ("Câmara de calçada", "Sim" if poco.get("camara_calcada") else "Não"),
            ("Responsável pela instalação", poco.get("responsavel_instalacao")),
        ]
        for indice in range(0, len(pares_poco), 2):
            for grupo, item in enumerate(pares_poco[indice : indice + 2]):
                coluna = grupo * 2
                rotulo, valor = item
                resumo.write(linha, coluna, rotulo, formato_rotulo)
                resumo.write(linha, coluna + 1, _normalizar_valor_excel(valor), formato_valor)
            linha += 1

    if desenvolvimento:
        linha += 1
        resumo.merge_range(linha, 0, linha, 3, "Desenvolvimento", formato_secao)
        linha += 1
        pares_desenv = [
            ("Realizado", "Sim" if desenvolvimento.get("realizado") else "Não"),
            ("Data", desenvolvimento.get("data")),
            ("Método", desenvolvimento.get("metodo")),
            ("Duração (min)", desenvolvimento.get("duracao_min")),
            ("NA antes (m)", desenvolvimento.get("na_antes_m")),
            ("NA depois (m)", desenvolvimento.get("na_depois_m")),
            ("Vazão (L/min)", desenvolvimento.get("vazao_l_min")),
            ("Volume retirado (L)", desenvolvimento.get("volume_retirado_l")),
            ("Turbidez inicial (NTU)", desenvolvimento.get("turbidez_inicial_ntu")),
            ("Turbidez final (NTU)", desenvolvimento.get("turbidez_final_ntu")),
        ]
        for indice in range(0, len(pares_desenv), 2):
            for grupo, item in enumerate(pares_desenv[indice : indice + 2]):
                coluna = grupo * 2
                rotulo, valor = item
                resumo.write(linha, coluna, rotulo, formato_rotulo)
                resumo.write(linha, coluna + 1, _normalizar_valor_excel(valor), formato_valor)
            linha += 1

    coluna_imagem = 5
    if imagem_perfil:
        resumo.insert_image(
            1,
            coluna_imagem,
            "perfil.png",
            {"image_data": io.BytesIO(imagem_perfil), "x_scale": 0.34, "y_scale": 0.34},
        )
    if imagem_construcao:
        resumo.insert_image(
            32,
            coluna_imagem,
            "construcao.png",
            {"image_data": io.BytesIO(imagem_construcao), "x_scale": 0.34, "y_scale": 0.34},
        )

    planilha_litologia = workbook.add_worksheet("Litologia")
    planilha_litologia.freeze_panes(1, 0)
    planilha_litologia.set_column("A:B", 14)
    planilha_litologia.set_column("C:C", 19)
    planilha_litologia.set_column("D:D", 55)
    planilha_litologia.set_column("E:F", 25)
    planilha_litologia.set_column("G:J", 15)
    _escrever_tabela_excel(
        planilha_litologia,
        workbook,
        camadas,
        [
            ("profundidade_inicial", "Prof. inicial (m)"),
            ("profundidade_final", "Prof. final (m)"),
            ("classificacao", "Classificação"),
            ("descricao_tatil_visual", "Descrição tátil-visual"),
            ("tipo_aquifero", "Unidade hidroestratigráfica"),
            ("zona_hidrica", "Condição hídrica"),
            ("cota_topo", "Cota do topo (m)"),
            ("cota_base", "Cota da base (m)"),
            ("espessura", "Espessura (m)"),
            ("espessura_vadosa", "Parcela vadosa (m)"),
        ],
    )

    planilha_pontos = workbook.add_worksheet("Amostras_VOC_NA")
    planilha_pontos.set_column("A:H", 22)
    linha_pontos = 0
    linha_pontos = _escrever_tabela_excel(
        planilha_pontos,
        workbook,
        coletas,
        [("profundidade_coleta", "Profundidade da coleta (m)")],
        linha_pontos,
    )
    linha_pontos = _escrever_tabela_excel(
        planilha_pontos,
        workbook,
        voc,
        [
            ("profundidade", "Profundidade VOC (m)"),
            ("concentracao", "Concentração (mg/L ou ppm)"),
        ],
        linha_pontos,
    )
    _escrever_tabela_excel(
        planilha_pontos,
        workbook,
        leituras_na,
        [
            ("data_hora", "Data e hora"),
            ("tipo", "Tipo de leitura"),
            ("profundidade_m", "Profundidade do NA (m)"),
            ("usar_como_estatico", "NA adotado"),
            ("observacoes", "Observações"),
        ],
        linha_pontos,
    )

    planilha_construcao = workbook.add_worksheet("Construcao_Poco")
    planilha_construcao.set_column("A:A", 24)
    planilha_construcao.set_column("B:C", 14)
    planilha_construcao.set_column("D:E", 28)
    planilha_construcao.set_column("F:H", 18)
    _escrever_tabela_excel(
        planilha_construcao,
        workbook,
        intervalos,
        [
            ("componente", "Componente"),
            ("profundidade_inicial", "Prof. inicial (m)"),
            ("profundidade_final", "Prof. final (m)"),
            ("material", "Material"),
            ("especificacao", "Especificação"),
            ("diametro_mm", "Diâmetro (mm)"),
            ("abertura_ranhura_mm", "Ranhura (mm)"),
            ("granulometria", "Granulometria"),
        ],
    )

    planilha_desenvolvimento = workbook.add_worksheet("Desenvolvimento")
    planilha_desenvolvimento.set_column("A:H", 20)
    desenvolvimento_df = pd.DataFrame([desenvolvimento]) if desenvolvimento else pd.DataFrame()
    linha_dev = _escrever_tabela_excel(
        planilha_desenvolvimento,
        workbook,
        desenvolvimento_df,
        [
            ("realizado", "Realizado"),
            ("data", "Data"),
            ("metodo", "Método"),
            ("duracao_min", "Duração (min)"),
            ("na_antes_m", "NA antes (m)"),
            ("na_depois_m", "NA depois (m)"),
            ("vazao_l_min", "Vazão (L/min)"),
            ("volume_retirado_l", "Volume retirado (L)"),
            ("turbidez_inicial_ntu", "Turbidez inicial (NTU)"),
            ("turbidez_final_ntu", "Turbidez final (NTU)"),
            ("ph_final", "pH final"),
            ("condutividade_final_us_cm", "Condutividade final (µS/cm)"),
            ("temperatura_final_c", "Temperatura final (°C)"),
            ("responsavel", "Responsável"),
            ("motivo_nao_realizado", "Motivo de não execução"),
            ("observacoes", "Observações"),
        ],
    )
    _escrever_tabela_excel(
        planilha_desenvolvimento,
        workbook,
        leituras_desenvolvimento,
        [
            ("tempo_min", "Tempo (min)"),
            ("nivel_agua_m", "NA (m)"),
            ("vazao_l_min", "Vazão (L/min)"),
            ("turbidez_ntu", "Turbidez (NTU)"),
            ("ph", "pH"),
            ("condutividade_us_cm", "Condutividade (µS/cm)"),
            ("temperatura_c", "Temperatura (°C)"),
            ("observacoes", "Observações"),
        ],
        linha_dev,
    )
    if imagem_desenvolvimento:
        planilha_desenvolvimento.set_column("J:R", 14)
        planilha_desenvolvimento.insert_image(
            1,
            9,
            "desenvolvimento.png",
            {
                "image_data": io.BytesIO(imagem_desenvolvimento),
                "x_scale": 0.45,
                "y_scale": 0.45,
                "object_position": 1,
            },
        )

    planilha_fotos = workbook.add_worksheet("Fotos")
    planilha_fotos.set_column("A:A", 8)
    planilha_fotos.set_column("B:B", 23)
    planilha_fotos.set_column("C:C", 15)
    planilha_fotos.set_column("D:D", 55)
    planilha_fotos.set_column("E:F", 24)
    cabecalhos_fotos = ["Nº", "Categoria", "Profundidade (m)", "Legenda", "Arquivo", "Data"]
    formato_cabecalho_fotos = workbook.add_format(
        {"bold": True, "font_color": "#FFFFFF", "bg_color": "#1F4E78", "border": 1}
    )
    formato_foto = workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"})
    for coluna, cabecalho in enumerate(cabecalhos_fotos):
        planilha_fotos.write(0, coluna, cabecalho, formato_cabecalho_fotos)
    planilha_fotos.freeze_panes(1, 0)

    if incluir_fotos:
        for indice, foto in enumerate(fotos[: max(0, int(limite_fotos))], start=1):
            linha_foto = indice
            planilha_fotos.set_row(linha_foto, 92)
            valores = [
                indice,
                foto.get("categoria"),
                foto.get("profundidade_m"),
                foto.get("legenda"),
                foto.get("nome_arquivo"),
                foto.get("criado_em"),
            ]
            for coluna, valor in enumerate(valores):
                planilha_fotos.write(
                    linha_foto,
                    coluna,
                    _normalizar_valor_excel(valor),
                    formato_foto,
                )
            try:
                miniatura = _preparar_foto_para_relatorio(foto["conteudo"], 800)
                planilha_fotos.insert_image(
                    linha_foto,
                    6,
                    f"foto_{indice}.jpg",
                    {
                        "image_data": miniatura,
                        "x_scale": 0.16,
                        "y_scale": 0.16,
                        "x_offset": 4,
                        "y_offset": 4,
                        "object_position": 1,
                    },
                )
            except Exception:
                pass
        planilha_fotos.set_column("G:G", 22)

    workbook.close()
    return memoria.getvalue()
